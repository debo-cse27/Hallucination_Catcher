import os
from docx import Document
from docx.shared import Pt, RGBColor

def create_word_report(pdf_filename, results_data, stats, output_dir="reports"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc = Document()
    
    # Title
    title = doc.add_heading('Automated Citation Verification Report', 0)
    
    # Document Info
    doc.add_paragraph(f"Source Document: {pdf_filename}")
    doc.add_paragraph(f"Total Citations Processed: {stats['Total']}")
    doc.add_paragraph(f"Hallucination Rate: {stats['Rate']}%")
    
    # Summary Statistics
    doc.add_heading('Summary Statistics', level=1)
    doc.add_paragraph(f"Verified Citations: {stats['Verified']}")
    doc.add_paragraph(f"Fabricated (Ghost) Citations: {stats['Ghost Citation']}")
    doc.add_paragraph(f"Mismatched Metadata: {stats['Mismatched Metadata']}")
    doc.add_paragraph(f"Errors / Incomplete: {stats['Other']}")
    
    # Detailed Analysis
    doc.add_heading('Detailed Citation Analysis', level=1)
    
    for i, item in enumerate(results_data, 1):
        ref = item['reference']
        ver = item['verification']
        status = ver['status']
        
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] Status: {status}")
        run.bold = True
        
        # Color Coding for library professionals to scan easily
        if status == "Verified":
            run.font.color.rgb = RGBColor(0, 128, 0) # Green
        elif status == "Ghost Citation":
            run.font.color.rgb = RGBColor(255, 0, 0) # Red
        elif status == "Mismatched Metadata":
            run.font.color.rgb = RGBColor(255, 140, 0) # Orange
            
        doc.add_paragraph(f"Claimed Title: {ref['title']}")
        if ref['doi']:
            doc.add_paragraph(f"Claimed DOI: {ref['doi']}")
            
        if status != "Verified":
            doc.add_paragraph(f"Registry Result: {ver['real_title']} (Match Score: {round(ver['score'])}%)")
            doc.add_paragraph("Action Required: Please manually review this citation.")
            
        doc.add_paragraph("-" * 40)
        
    report_path = os.path.join(output_dir, f"{pdf_filename.replace('.pdf', '')}_Report.docx")
    doc.save(report_path)
    return report_path